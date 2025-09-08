import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Form, status, Request, BackgroundTasks
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
import asyncio
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
import openai
from openai import AzureOpenAI
import json
import base64
import os
import pandas as pd
import tempfile
import html
import traceback
import re
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from docx import Document
import io
from bs4 import BeautifulSoup, NavigableString
import hashlib
from datetime import datetime, timezone, timedelta
from pymongo import MongoClient
from opencensus.ext.azure.trace_exporter import AzureExporter
from opencensus.trace.samplers import ProbabilitySampler
from opencensus.trace.tracer import Tracer
from azure.core.credentials import AzureKeyCredential
from azure.ai.contentsafety import ContentSafetyClient
from azure.ai.contentsafety.models import AnalyzeTextOptions, TextCategory
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.models import VectorizedQuery
from langchain_openai import AzureOpenAIEmbeddings
from langchain.schema import HumanMessage
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
import time
import logging

# --- SECURITY IMPORTS ---
from jose import JWTError, jwt
from passlib.context import CryptContext
import jwt as pyjwt # Using pyjwt for external token validation
from jwt.algorithms import RSAAlgorithm

# --- OBSERVABILITY IMPORTS FOR ARIZE PHOENIX ---
from opentelemetry import trace as trace_api
from opentelemetry.exporter.otlp.proto.http.trace_exporter import OTLPSpanExporter
from opentelemetry.sdk.resources import Resource
from opentelemetry.sdk.trace import TracerProvider
from opentelemetry.sdk.trace.export import SimpleSpanProcessor
from openinference.instrumentation.openai import OpenAIInstrumentor

# --- Load Environment Variables ---
load_dotenv()

# --- Professional Logging Setup ---
# For production, level=logging.INFO is standard to keep logs clean.
# For development, change to logging.DEBUG to see all diagnostic messages.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- MODIFICATION: Silence noisy SDK loggers in production ---
# These loggers produce verbose request/response details at the INFO level.
# By setting them to WARNING, we only see them if something is wrong.
logging.getLogger("azure.core.pipeline.policies").setLevel(logging.WARNING)
logging.getLogger("openai").setLevel(logging.WARNING)


# --- SECURITY CONFIGURATION ---
SECRET_KEY = os.environ.get("SECRET_KEY", "a_very_secret_key_that_should_be_in_env_for_production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60
DEV_MODE = os.environ.get("DEV_MODE", "false").lower() == "true"

# --- Entra ID / SSO Configuration ---
TENANT_ID = os.environ.get("TENANT_ID")
CLIENT_ID = os.environ.get("CLIENT_ID")
ISSUER_URL = f"https://login.microsoftonline.com/{TENANT_ID}/v2.0"
JWKS_URL = f"https://login.microsoftonline.com/{TENANT_ID}/discovery/v2.0/keys"


# --- ARIZE PHOENIX CONFIGURATION ---
PHOENIX_COLLECTOR_ENDPOINT = os.environ.get("PHOENIX_COLLECTOR_ENDPOINT")
PHOENIX_API_KEY = os.environ.get("PHOENIX_API_KEY")

# --- Initialize Phoenix Tracer ---
def initialize_phoenix_tracer():
    if not PHOENIX_COLLECTOR_ENDPOINT:
        logging.info("Phoenix collector endpoint not configured. Skipping Phoenix initialization.")
        return

    headers = {"Authorization": f"Bearer {PHOENIX_API_KEY}"} if PHOENIX_API_KEY else {}
    resource = Resource(attributes={"service.name": "bial-regulatory-platform"})
    
    trace_provider = TracerProvider(resource=resource)
    span_exporter = OTLPSpanExporter(endpoint=PHOENIX_COLLECTOR_ENDPOINT, headers=headers)
    trace_provider.add_span_processor(SimpleSpanProcessor(span_exporter))
    trace_api.set_tracer_provider(trace_provider)

    # Auto-instrument key libraries
    OpenAIInstrumentor().instrument()
    logging.info("Arize Phoenix tracer initialized and instrumentation applied.")

initialize_phoenix_tracer()


pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# --- Helper Function for Credentials ---
def check_creds(cred_value, placeholder_prefix="YOUR_"):
    """Checks if a credential is a placeholder or missing."""
    if not cred_value or placeholder_prefix in cred_value.upper():
        return True
    return False

# --- Configuration ---
class AzureCredentials(BaseModel):
    search_endpoint: str = os.environ.get("SEARCH_ENDPOINT")
    search_api_key: str = os.environ.get("SEARCH_API_KEY")
    default_search_index_name: str = os.environ.get("DEFAULT_SEARCH_INDEX_NAME")
    default_vector_field_name: str = os.environ.get("DEFAULT_VECTOR_FIELD_NAME")
    default_semantic_config_name: str = os.environ.get("DEFAULT_SEMANTIC_CONFIG_NAME")
    openai_endpoint: str = os.environ.get("OPENAI_ENDPOINT")
    openai_api_version: str = os.environ.get("OPENAI_API_VERSION")
    openai_api_key: str = os.environ.get("OPENAI_API_KEY")
    deployment_id: str = os.environ.get("DEPLOYMENT_ID")
    planning_llm_deployment_id: str = os.environ.get("PLANNING_LLM_DEPLOYMENT_ID")
    embedding_deployment_id: str = os.environ.get("EMBEDDING_DEPLOYMENT_ID")
    cosmos_mongo_connection_string: str = os.environ.get("COSMOS_MONGO_CONNECTION_STRING")
    cosmos_database_name: str = os.environ.get("COSMOS_DATABASE_NAME")
    cosmos_logs_collection: str = os.environ.get("COSMOS_LOGS_COLLECTION")
    cosmos_users_collection: str = os.environ.get("COSMOS_USERS_COLLECTION", "Users")
    cosmos_feedback_collection: str = os.environ.get("COSMOS_FEEDBACK_COLLECTION", "Feedback")
    app_insights_connection_string: str = os.environ.get("APP_INSIGHTS_CONNECTION_STRING")
    content_safety_endpoint: str = os.environ.get("CONTENT_SAFETY_ENDPOINT")
    content_safety_key: str = os.environ.get("CONTENT_SAFETY_KEY")
    bing_search_api_key: str = os.environ.get("BING_SEARCH_API_KEY")
    bing_search_endpoint: str = os.environ.get("BING_SEARCH_ENDPOINT", "https://api.bing.microsoft.com/v7.0/search")
    serpapi_api_key: str = os.environ.get("SERPAPI_API_KEY")

config = AzureCredentials()
app = FastAPI(title="BIAL Regulatory Platform API")

# --- CORS Middleware ---
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"]
)

# --- AZURE MONITOR USER TRACKING MIDDLEWARE ---
@app.middleware("http")
async def add_user_to_trace(request: Request, call_next):
    response = await call_next(request)
    try:
        if "current_user" in request.state:
            user = request.state.current_user
            tracer = trace_api.get_tracer(__name__)
            span = trace_api.get_current_span()
            span.set_attribute("enduser.id", user.username)
    except Exception:
        pass
    return response


# --- Service Clients (Graceful Initialization) ---
class ServiceClients:
    def __init__(self):
        # OpenAI Clients
        if not check_creds(config.openai_api_key) and not check_creds(config.openai_endpoint):
            self.synthesis_openai_client = AzureOpenAI(api_key=config.openai_api_key, azure_endpoint=config.openai_endpoint, api_version=config.openai_api_version)
            self.planning_openai_client = self.synthesis_openai_client
            self.search_query_embeddings_model = AzureOpenAIEmbeddings(azure_deployment=config.embedding_deployment_id, azure_endpoint=config.openai_endpoint, api_key=config.openai_api_key, api_version=config.openai_api_version, chunk_size=1)
        else:
            self.synthesis_openai_client = None
            self.planning_openai_client = None
            self.search_query_embeddings_model = None

        # Content Safety Client
        if not check_creds(config.content_safety_endpoint) and not check_creds(config.content_safety_key):
            self.content_safety_client = ContentSafetyClient(config.content_safety_endpoint, AzureKeyCredential(config.content_safety_key))
        else:
            self.content_safety_client = None

        # Mongo Client
        if not check_creds(config.cosmos_mongo_connection_string):
            self.mongo_client = MongoClient(config.cosmos_mongo_connection_string)
        else:
            self.mongo_client = None
        
        # Tracer
        if not check_creds(config.app_insights_connection_string):
            self.tracer = Tracer(exporter=AzureExporter(connection_string=config.app_insights_connection_string), sampler=ProbabilitySampler(1.0))
        else:
            self.tracer = None
            
        # Fetch and cache Microsoft's signing keys for SSO
        try:
            self.jwks_client = pyjwt.PyJWKClient(JWKS_URL)
        except Exception as e:
            logging.error(f"Could not fetch JWKS from Microsoft: {e}")
            self.jwks_client = None

clients = ServiceClients()

# --- Pydantic Models ---
class ChatMessage(BaseModel):
    role: str
    content: str

class ConversationalChatResponse(BaseModel):
    answer: str
    plan: List[str]
    sources: List[Dict[str, Any]]
    source: str

class ChatRequest(BaseModel):
    question: str
    history: List[ChatMessage] = []

class AnalysisResponse(BaseModel):
    report: str

class RefineReportRequest(BaseModel):
    original_report: str
    new_info: str

class RefineReportResponse(BaseModel):
    refined_report: str
    
class DownloadRequest(BaseModel):
    html_content: str

class SSOLoginRequest(BaseModel):
    sso_token: str

class User(BaseModel):
    username: str

class UserInDB(User):
    hashed_password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class TokenData(BaseModel):
    username: Optional[str] = None

class UserCreate(BaseModel):
    username: str
    password: str

class FeedbackRequest(BaseModel):
    question: str
    answer: str
    feedback: str # "like" or "dislike"

# --- Shared System Prompt ---
SHARED_SYSTEM_PROMPT = """.Part 1:   "you are AI assistant for Multiyear tarrif submission for AERA. final response should have 1500 words at least.
* **3.2. Source Attribution (Authority Data - Handling Terminology for Authority's Stance):**
    If the query relates to the "Authority's" stance (e.g., user asks “What is the authority’s *approved* change?”, “What did the authority *decide*?”, “What was *proposed* by the authority?”), your primary goal is to find the authority's documented action or position on that specific subject *within the  .
    * **Understanding Levels of Finality (for AI's internal logic):**
        * **Final/Conclusive Terms:** "approved by the authority", "decided by the authority", "authority's decision", "final tariff/order states", "sanctioned by authority", "adopted by authority".
        * **Provisional/Draft Terms:** "proposed by theauthority", "authority's proposal", "draft figures", "recommended by the authority" (if it's a recommendation for a later decision).
        * **Analytical/Consideration Terms:** "considered by the authority", "analyzed by the authority", "examined by the authority", "authority's review/view/assessment/preliminary findings".
    * **Extraction Strategy Based on User Intent and Document Content:**
        1.  **Attempt to Match User's Exact Term First:** Always search the CONTEXT for information explicitly matching the user's specific terminology (e.g., if the user asks for "approved," look first for "approved by the authority"). If found, present this.
        2.  **If User's Query Implies Finality (e.g., asks for "approved," "final figures," "decision"):**
            * And their *exact term* is NOT found in the CONTEXT for that item:
                * **Prioritize searching the CONTEXT for other Final/Conclusive terms** (e.g., "decided by the authority," "authority's decision"). If one of these is found, present this information. You MUST then state clearly: "You asked for 'approved.' The document describes what was '*[actual term found, e.g., decided by the authority]*' as follows: \[data and references]."
                * If no Final/Conclusive terms are found for that item in the CONTEXT, then (and only then) look for Provisional/Draft terms (e.g., "proposed by the authority"). If found, present this, stating: "You asked for 'approved.' A final approval or decision was not found for this item in the provided context. However, the authority '*proposed*' the following: \[data and references]."
                * If neither of the above is found, look for Analytical/Consideration terms and report similarly with clarification.
        3.  **If User's Query Uses a Provisional/Draft Term (e.g., "proposed"):**
            * Prioritize finding information matching those Provisional/Draft terms in the CONTEXT.
            * If not found, you can then look for Analytical/Consideration terms, clarifying the terminology. Avoid presenting Final/Conclusive terms unless you explicitly state that the user asked for a draft but a final version was found regarding that specific point.
        4.  **If User's Query Uses an Analytical/Consideration Term (e.g., "considered"):**
            * Prioritize finding information matching those terms. Always give me the table refrence for the response i mean which table no you have refered for the response
    * **Accurate Reporting is Key:** Always present information using the **document's actual terminology**. Clearly explain if and how it relates to the user's original query terms. If multiple relevant stages of authority action are evident in the CONTEXT (e.g., a proposal and then a later decision), you may summarize both, clearly distinguishing them by the terms used in the document.
    * **If No Relevant Authority Action Found:** If the provided CONTEXT contains no clear information matching any relevant stage of authority action (final, provisional, or analytical) regarding the specific subject of the query, state that this information was not found for the authority in the provided context.
    * **Table Headers:** Use data from tables if their headers clearly indicate the source and nature of the data (e.g., "Figures as Decided by the Authority," "Operator's Proposed Traffic").
"**Crucial Instruction for Authority's Stance:** When the user asks for 'approved', 'final', or 'decided' figures from the Authority, it is **imperative** that you prioritize extracting text explicitly labeled with terms like 'decided by the authority' or 'approved by the authority' from the CONTEXT. If such conclusive terms are present for the queried item, present them. Only if NO such conclusive terms are found in the CONTEXT for that item should you then present information labeled 'proposed by the authority', and you MUST clearly state that you are providing 'proposed' figures because 'approved/decided' ones were not found in the given context."
* **Clarifying "Considered by Authority" in a "True-Up" Context:**
    * If the user query asks what the authority 'considered' in relation to a 'true-up' of a specific control period:
        * First, check if the CONTEXT contains information about the authority's **analysis or verification of the actual figures submitted for that true-up period**. If so, present this.
        * If the user's query might also imply understanding the **original baseline** that is being trued-up against, you can additionally (or if the above is not found) look for what the authority **originally considered or determined when setting the tariffs for that control period at the beginning of that period**.
        * **Crucially, always differentiate these two.** For example: "For the true-up of the Third Control Period, DIAL submitted the following actual traffic figures (e.g., from Table 25): [data]. The figures that the Authority had originally considered at the time of determining the tariff for the Third Control Period were (e.g., from Table 26): [data]."
        * If the query is simply "What was considered for true-up..." without specifying "original determination" vs "actuals review", and both types of information are in the context, you might offer both or ask the user to clarify which aspect of "considered for true-up" they are interested in Mandatory Table Referencing:
For any data, figures, or claims extracted from a table within the CONTEXT, you must cite the corresponding table number in your response. This is a strict requirement for all outputs. The reference should be placed directly with the data it pertains to.
Example 1: "The Authority approved Aeronautical Revenue of ₹1,500 Cr for FY 2024-25 (Table 15)."
Example 2: "For the true-up, the operator submitted actual passenger traffic of 45 million (as per Table 3.2), while the original figure considered by the Authority was 42 million (from Table 5.1 of the original Order)
* **Mandatory Table Referencing:**
For any data, figures, or claims extracted from a table within the CONTEXT, you must cite the corresponding table number in your response. This is a strict requirement for all outputs. The reference should be placed directly with the data it pertains to.
Example: "The Authority approved Aeronautical Revenue of ₹1,500 Cr for FY 2024-25 (Table 15)."
PECIFIC INSTRUCTIONS FOR YOUR RESPONSE (in addition to the general background provided):
1. Directly address all parts of the ORIGINAL USER QUESTION.
2. Synthesize information from the different context sections if they relate to different aspects of the original question.
3. Format numerical data extracted from tables into an HTML table with borders (e.g., <table border='1'>...). Use table headers (<th>) and table data cells (<td>).
4. **References are crucial.** At the end of your answer, include a 'References:' section listing the source documents.
*Crucially:* Include references for the information presented. Mention the specific source and, if mentioned within the text context itself, include table numbers (e.g., 'Table 26') or section titles and the file name. Present these references clearly at the end of your answer under a 'References:' heading.

"""

# --- AUTHENTICATION HELPERS ---
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def get_user_from_db(username: str):
    if not clients.mongo_client: return None
    user_collection = clients.mongo_client[config.cosmos_database_name][config.cosmos_users_collection]
    user_doc = user_collection.find_one({"username": username})
    if user_doc:
        return UserInDB(username=user_doc["username"], hashed_password=user_doc["password_hash"])
    return None

async def get_current_user_prod(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None: raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
    return User(username=token_data.username)

async def get_current_user_dev_and_set_state(request: Request):
    user = User(username="dev_user")
    request.state.current_user = user
    return user

async def get_current_user_and_set_state(request: Request, token: str = Depends(oauth2_scheme)):
    user = await get_current_user_prod(token)
    request.state.current_user = user
    return user

auth_dependency = get_current_user_dev_and_set_state if DEV_MODE else get_current_user_and_set_state

# --- LOGGING FUNCTION ---
def log_interaction(user: str, question: str, answer: str, duration: float):
    if not clients.mongo_client:
        logging.warning("Could not log interaction: Database service is not configured.")
        return
    try:
        log_collection = clients.mongo_client[config.cosmos_database_name][config.cosmos_logs_collection]
        log_collection.insert_one({
            "timestamp": datetime.now(timezone.utc), 
            "user": user, 
            "question": question, 
            "answer": answer, 
            "duration": duration
        })
    except Exception as e:
        logging.error(f"Failed to write to custom audit log: {e}")


# --- Core Logic Functions ---
def extract_text_from_docx(file: UploadFile) -> str:
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file.file.read())
            tmp_path = tmp.name
        document = docx.Document(tmp_path)
        text = "\n".join([para.text for para in document.paragraphs])
        os.remove(tmp_path)
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading Word document: {e}")

def query_azure_search(query_text: str, index_name: str, k: int = 5):
    if check_creds(config.search_endpoint) or check_creds(config.search_api_key):
        return "Error: Azure Search service is not configured.", []
    try:
        search_client = SearchClient(config.search_endpoint, index_name, AzureKeyCredential(config.search_api_key))
        
        search_kwargs = {
            "search_text": query_text if query_text and query_text.strip() else "*",
            "top": k,
            "query_type": "semantic",
            "semantic_configuration_name": config.default_semantic_config_name,
            "query_caption": "extractive",
            "query_answer": "extractive"
        }
        
        results = search_client.search(**search_kwargs)
        
        context = "\n\n".join([doc.get("content", "") for doc in results])
        references_data = [{"filename_or_title": doc.get("title") or os.path.basename(doc.get("filepath", "")), "url": doc.get("url")} for doc in results]
        return context.strip(), references_data
    except Exception as e:
        return f"Error accessing search index: {e}", []

def query_serpapi(query: str, count: int = 5) -> str:
    if check_creds(config.serpapi_api_key): return "Error: SerpApi API key not configured."
    params = {"q": query, "api_key": config.serpapi_api_key, "num": count}
    try:
        response = requests.get("https://serpapi.com/search.json", params=params, timeout=15)
        response.raise_for_status()
        search_results = response.json()
        organic_results = search_results.get("organic_results", [])
        snippets = [f"Title: {res.get('title', 'N/A')}\nSnippet: {res.get('snippet', 'N/A')}" for res in organic_results]
        return "\n".join(snippets) if snippets else "No web search results found via SerpApi."
    except Exception as e: return f"Error during SerpApi search: {e}"

def query_bing_web_search(query: str, count: int = 5) -> str:
    if check_creds(config.bing_search_api_key): return "Error: Bing Search API key not configured."
    headers = {"Ocp-Apim-Subscription-Key": config.bing_search_api_key}
    params = {"q": query, "count": count}
    try:
        response = requests.get(config.bing_search_endpoint, headers=headers, params=params, timeout=10)
        response.raise_for_status()
        search_results = response.json()
        snippets = [f"Title: {res['name']}\nSnippet: {res['snippet']}" for res in search_results.get("webPages", {},).get("value", [])]
        return "\n".join(snippets) if snippets else "No web search results found."
    except Exception as e: return f"Error during Bing web search: {e}"

def get_query_plan_from_llm(user_question: str, history: List[ChatMessage]):
    history_str = "\n".join([f'{msg.role}: {msg.content}' for msg in history])
    planning_prompt = f"""You are a query planning assistant specializing in breaking down complex questions about **AERA regulatory documents, often concerning tariff orders, consultation papers, control periods, and specific financial data (like CAPEX, Opex, Traffic) for airport operators such as DIAL, MIAL, BIAL, HIAL.**
Your primary task is to take a user's complex question related to these topics and break it down into a series of 1 to 20 simple, self-contained search queries that can be individually executed against a document index. Each search query should aim to find a specific piece of information (e.g., a specific figure, a justification, a comparison point) needed to answer the overall complex question.
If the user's question is already simple and can be answered with a single search, return just that single query in the list.
If the question is very complex and might require more distinct search steps, formulate the most critical 1 to 20 search queries, focusing on distinct pieces of information.
Return your response ONLY as a JSON list of strings, where each string is a search query.

CONVERSATION HISTORY:
{history_str}

User's complex question: {user_question}
Your JSON list of search queries:"""
    try:
        response = clients.planning_openai_client.chat.completions.create(model=config.planning_llm_deployment_id, messages=[{"role": "user", "content": planning_prompt}], temperature=0.0, max_tokens=16000)

        plan_str = response.choices[0].message.content
        if match := re.search(r'\[.*\]', plan_str, re.DOTALL):
            return json.loads(match.group(0))
        return [user_question]
    except Exception as e:
        logging.error(f"Error in query planner: {e}")
        return [user_question]

async def execute_advanced_rag_pipeline(user_question: str, history: Optional[List[ChatMessage]] = None):
    """
    Executes the full advanced RAG pipeline: Plan -> Retrieve (multi-query) -> Synthesize.
    This is used by both conversational chat and document analysis.
    """
    if not clients.synthesis_openai_client or not clients.planning_openai_client:
        raise HTTPException(status_code=503, detail="AI services are not configured.")

    if history is None:
        history = []

    # 1. Get query plan from the planning LLM
    logging.info(f"Generating query plan for: '{user_question[:100]}...'")
    query_plan = get_query_plan_from_llm(user_question, history)
    
    # 2. Execute the plan: retrieve context for each sub-query
    logging.info(f"Executing plan with {len(query_plan)} steps...")
    combined_context = ""
    all_retrieved_details = []
    for sub_query in query_plan:
        logging.debug(f"  - Searching for: '{sub_query}'")
        context_for_step, retrieved_details_for_step = query_azure_search(sub_query, config.default_search_index_name, k=5)
        if context_for_step and not context_for_step.startswith("Error"):
            combined_context += f"\n\n--- Context for sub-query: '{sub_query}' ---\n" + context_for_step
            all_retrieved_details.extend(retrieved_details_for_step)

    # De-duplicate sources based on URL or title
    unique_sources = list({(item.get('url') or item.get('filename_or_title')): item for item in all_retrieved_details}.values())

    # 3. Handle cases with no context
    if not combined_context.strip():
        logging.warning("No relevant information found in search index.")
        return {
            "answer": "I could not find any relevant information in the internal documents to answer your question.",
            "plan": query_plan,
            "sources": unique_sources
        }

    # 4. Synthesize the final answer using the aggregated context
    logging.info("Synthesizing final answer from aggregated context...")
    
    sources_text = "\n".join([f"- {source.get('filename_or_title', 'Unknown Source')}" for source in unique_sources])

    synthesis_prompt = f"""Based on the following aggregated context from multiple search steps, please synthesize a comprehensive answer to the original user question.

ORIGINAL USER QUESTION / INSTRUCTION:
{user_question}

AGGREGATED CONTEXT:
---------------------
{combined_context}
---------------------

IDENTIFIED SOURCES:
---------------------
{sources_text}
---------------------

SPECIFIC INSTRUCTIONS FOR YOUR RESPONSE (in addition to the general background provided):\n"
            1. Directly address all parts of the ORIGINAL USER QUESTION.\n"
            2. Synthesize information from the different context sections if they relate to different aspects of the original question.\n"
            3. Format numerical data extracted from tables into an HTML table with borders (e.g., <table border='1'>...). Use table headers (<th>) and table data cells (<td>).\n"
            4. **References are crucial.** At the end of your answer, include a 'References:' section listing the source documents (using filenames or titles as provided in 'IDENTIFIED CONTEXT SOURCES') from which the information was derived. If a URL is available for a source, make the filename/title a clickable hyperlink to that URL.\n\n"
            COMPREHENSIVE ANSWER TO THE ORIGINAL USER QUESTION:\n"
            *Crucially:* Include references for the information presented. Mention the specific source (e.g., the filename from the IDENTIFIED CONTEXT SOURCES list) and, if mentioned within the text context itself, include table numbers (e.g., 'Table 26') or section titles and the file name. Present these references clearly at the end of your answer under a 'References:' heading."
            "\n\nANSWER:"
            

YOUR COMPREHENSIVE ANSWER:"""
    
    messages = [
        {"role": "system", "content": SHARED_SYSTEM_PROMPT},
        *[msg.dict() for msg in history],
        {"role": "user", "content": synthesis_prompt}
    ]
    
    response = clients.synthesis_openai_client.chat.completions.create(
        model=config.deployment_id, 
        messages=messages, 
        temperature=0.2, 
        max_tokens=16000
    )
    final_answer = response.choices[0].message.content

    return {
        "answer": final_answer,
        "plan": query_plan,
        "sources": unique_sources
    }

# --- CORRECTED: Robust Chat Request Handler with Safe Logging ---
async def handle_chat_request(request: ChatRequest, current_user: User):
    """
    Handles all chat requests with improved error handling and logging.
    """
    start_time = time.time()
    try:
        if not clients.synthesis_openai_client:
            raise HTTPException(status_code=503, detail="Synthesis LLM not configured.")
        
        web_search_keywords = ["web search", "latest", "current", "internet"]
        is_web_search = any(w in request.question.lower() for w in web_search_keywords)
        
        response_data = None
        if is_web_search:
            source = "web"
            logging.info(f"Web search triggered for: '{request.question}'")
            
            combined_context = ""
            if not check_creds(config.serpapi_api_key):
                logging.info("Attempting web search with SerpApi...")
                combined_context = query_serpapi(request.question)

            if not combined_context or combined_context.startswith("Error"):
                if not check_creds(config.bing_search_api_key):
                    logging.info("SerpApi failed or not configured, attempting web search with Bing Search...")
                    combined_context = query_bing_web_search(request.question)
                else:
                    if not combined_context:
                        combined_context = "Error: No web search API key configured."
            
            if not combined_context.strip() or combined_context.startswith("Error"):
                 response_data = ConversationalChatResponse(answer="I could not retrieve any information from the web search.", plan=[request.question], sources=[], source=source)
            else:
                messages = [
                    {"role": "system", "content": SHARED_SYSTEM_PROMPT},
                    *[msg.dict() for msg in request.history],
                    {"role": "user", "content": f"Based on the following context from a web search, answer the user's question.\n\nUser Question: {request.question}\n\nCONTEXT:\n{combined_context}"}
                ]
                response = clients.synthesis_openai_client.chat.completions.create(model=config.deployment_id, messages=messages, temperature=0.2, max_tokens=16000)
                final_answer = response.choices[0].message.content
                response_data = ConversationalChatResponse(answer=final_answer, plan=[request.question], sources=[], source=source)
        else:
            source = "internal"
            logging.info(f"Internal RAG triggered for: '{request.question}'")
            rag_result = await execute_advanced_rag_pipeline(request.question, request.history)
            response_data = ConversationalChatResponse(
                answer=rag_result["answer"],
                plan=rag_result["plan"],
                sources=rag_result["sources"],
                source=source
            )

        # Log successful interaction before returning
        duration = time.time() - start_time
        log_interaction(user=current_user.username, question=request.question, answer=response_data.answer, duration=duration)
        return response_data

    except Exception as e:
        # Log the error and then re-raise the HTTPException
        duration = time.time() - start_time
        error_message = f"Error: {str(e)}"
        logging.error(f"AN ERROR OCCURRED IN handle_chat_request: {e}", exc_info=True)
        log_interaction(user=current_user.username, question=request.question, answer=error_message, duration=duration)
        raise HTTPException(status_code=500, detail=f"Error generating answer: {e}")


# --- CORRECTED: Comprehensive HTML to DOCX Parsing Function ---
def parse_html_to_docx(soup, document):
    """
    Parses a BeautifulSoup object and adds the content to a python-docx Document.
    This version iterates through all children of the body to capture raw text nodes
    that are not wrapped in <p> tags, ensuring all text is included.
    """
    # Determine the root element to iterate over. Fallback to soup if body doesn't exist.
    content_root = soup.body if soup.body else soup
    
    for element in content_root.children:
        # If the element is a raw string of text (NavigableString)
        if isinstance(element, NavigableString):
            if element.string and element.string.strip():
                document.add_paragraph(element.string.strip())
            continue

        # If the element is a tag
        if hasattr(element, 'name') and element.name:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                try:
                    level = int(element.name[1])
                    document.add_heading(element.get_text(strip=True), level=level)
                except (ValueError, IndexError):
                    document.add_heading(element.get_text(strip=True), level=2)
            
            elif element.name == 'p':
                p = document.add_paragraph()
                for content in element.contents:
                    if hasattr(content, 'name') and content.name in ['b', 'strong']:
                        p.add_run(content.get_text(strip=True)).bold = True
                    elif hasattr(content, 'name') and content.name in ['i', 'em']:
                        p.add_run(content.get_text(strip=True)).italic = True
                    else:
                         # Use str() to handle NavigableStrings and other tag contents
                        p.add_run(str(content))

            elif element.name in ['ul', 'ol']:
                style = 'List Bullet' if element.name == 'ul' else 'List Number'
                for li in element.find_all('li', recursive=False):
                    document.add_paragraph(li.get_text(strip=True), style=style)

            elif element.name == 'table':
                rows = element.find_all('tr')
                if not rows: continue
                
                # More robustly find header and body rows
                header_rows = element.select('thead > tr')
                body_rows = element.select('tbody > tr')
                
                if not header_rows and not body_rows: # Simple table with no thead/tbody
                    header_rows = rows[0:1]
                    body_rows = rows[1:]
                elif not body_rows: # Table might only have a thead
                    body_rows = []
                
                header_cells = header_rows[0].find_all(['th', 'td']) if header_rows else []
                if not header_cells: continue
                
                table = document.add_table(rows=0, cols=len(header_cells))
                table.style = 'Table Grid'
                table.autofit = True

                for h_row in header_rows:
                    cells = h_row.find_all(['th', 'td'])
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(cells):
                        if i < len(row_cells):
                            p = row_cells[i].paragraphs[0]
                            p.add_run(cell.get_text(strip=True)).bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                for b_row in body_rows:
                    cells = b_row.find_all('td')
                    if len(cells) == len(header_cells): # Ensure row has correct number of cells
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                           row_cells[i].text = cell.get_text(strip=True)

            elif element.name == 'hr':
                document.add_page_break()


# --- API Endpoints ---
@app.get("/")
def read_root():
    return {"message": "BIAL Regulatory Platform API is running."}

@app.post("/register", response_model=User)
async def register_user(user: UserCreate):
    if not clients.mongo_client:
        raise HTTPException(status_code=503, detail="Database service is not configured.")
    user_collection = clients.mongo_client[config.cosmos_database_name][config.cosmos_users_collection]
    if user_collection.find_one({"username": user.username}):
        raise HTTPException(status_code=400, detail="Username already registered")
    
    hashed_password = get_password_hash(user.password)
    user_collection.insert_one({"username": user.username, "password_hash": hashed_password})
    return User(username=user.username)

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    if not clients.mongo_client:
        raise HTTPException(status_code=503, detail="Database service is not configured.")
    user = get_user_from_db(form_data.username)
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(data={"sub": user.username}, expires_delta=access_token_expires)
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/login/sso", response_model=Token)
async def login_sso(request: SSOLoginRequest):
    if not all([TENANT_ID, CLIENT_ID]):
        raise HTTPException(status_code=503, detail="SSO environment variables not configured on server.")
    if not clients.jwks_client:
        raise HTTPException(status_code=503, detail="SSO service is not configured correctly (JWKS keys not found).")
    try:
        signing_key = clients.jwks_client.get_signing_key_from_jwt(request.sso_token).key
        
        decoded_token = pyjwt.decode(
            request.sso_token,
            signing_key,
            algorithms=["RS256"],
            audience=CLIENT_ID,
            issuer=ISSUER_URL
        )
        
        username = decoded_token.get("preferred_username") or decoded_token.get("upn")
        if not username:
            raise HTTPException(status_code=400, detail="Username not found in SSO token.")

        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        internal_access_token = create_access_token(
            data={"sub": username}, expires_delta=access_token_expires
        )
        return {"access_token": internal_access_token, "token_type": "bearer"}

    except pyjwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="SSO token has expired.")
    except pyjwt.InvalidAudienceError:
        raise HTTPException(status_code=401, detail="Invalid SSO token audience.")
    except pyjwt.InvalidIssuerError:
        raise HTTPException(status_code=401, detail="Invalid SSO token issuer.")
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Could not validate SSO token: {e}")

@app.post("/analyze-document", response_model=AnalysisResponse)
async def analyze_document(analysis_title: str = Form(...), file: UploadFile = File(...), current_user: User = Depends(auth_dependency)):
    """
    Analyzes an uploaded document using the unified Advanced RAG pipeline.
    This endpoint now benefits from query planning for more thorough analysis.
    """
    start_time = time.time()
    final_report = ""
    try:
        if not clients.synthesis_openai_client:
            raise HTTPException(status_code=503, detail="OpenAI service is not configured.")
    
        extracted_text = extract_text_from_docx(file)
        summary_prompt = f"Please provide a detailed, neutral summary of the key points of the following document text:\n\n---\n{extracted_text[:20000]}\n---"
        document_summary = clients.synthesis_openai_client.chat.completions.create(model=config.deployment_id, messages=[{"role": "user", "content": summary_prompt}]).choices[0].message.content

        analysis_prompts_config = {
            "MDA Manpower Analysis": {
           "Analysis of manpower expenditure projection for BIAL for fourth control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the manpower expenditure projected by BIAL: 1- Year on Year growth of personnel cost projected by BIAL for fourth control period. 2- Justification for personnel cost growth in fourth control period provided by BIAL. 3- year on year Manpower Expenses growth Submitted by DIAL for fourth control period in DIAL fourth control period consultation Paper. 4- Justification provided by DIAL for manpower expenses submitted by DIAL for fourth control period. 5- Examination and rationale provided by authority for manpower expenses submitted by DIAL for fourth control period. 6- Year on Year growth of employee cost submitted by MIAL for fourth control period for fourth control period in MIAL Fourth control consultation Paper. 7- Justification provided by MIAL for manpower expenses per passeneger traffic submitted by MIAL for fourth control period. 8- Examination and rationale provided by authority for manpower expenses submitted by MIAL for fourth control period. 9- Using the rationale extracted in steps 4, 5 7 and 8 suggest how the rationale or justification provided by BIAL in the MDA document for manpower expenditure for fourth control period can be enhanced. For every suggestion made, give specific reason why the suggestion was made by you using relevant references from DIAL and MIAL tariff orders or consultation papers.",
           "Analysis of actual manpower expenditure for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analyzing the actual manpower expenditure for the third control period: 1. Actual manpower expenditure for BIAL and variance from authority approved manpower expenditure for the third control period. 2. Justification for manpower expenditure in third control period provided by BIAL. 3 Actual manpower expenditure for DIAL and variance from authority approved manpower expenditure for the third control period. 4. Justification provided by DIAL for actual manpower expenses for third control period and the reason for variance compared to authority approved figures. 5. Examination and rationale provided by authority for actual manpower expenditure for only DIAL for third control period and its variance compared to authority approved figures. 6. Actual manpower expenditure for MIAL and variance from authority approved manpower expenditure for the third control period. 7. Justification provided by MIAL for actual manpower expenses submitted by MIAL for third control period and the reason for variance with authority approved figures. 8. Examination and rationale provided by authority for actual manpower expenditure submitted by only MIAL for third control period and its variance compared to authority approved figures. 9. Using the rationale extracted in steps 4, 5, 7, and 8, suggest how the rationale or justification provided by BIAL in the MDA document for manpower expenditure for the third control period can be enhanced. For every suggestion made, give specific reason why the suggestion was made by you using relevant references from DIAL and MIAL tariff orders or consultation papers.",
           "Analysis of KPI Computation for BIAL for fourth Control period": f"the upload document proposes the following: '{{document_summary}}'. Use the following steps for analyzing the KPI Computation.Calculate and compare the YoY change of employee expenses of DIAL and MIAL for the fourth control period,first give what is total manpower expense submitted by DIAL for fourth control period , employee cost submitted by MIAL for fourth control period . after wards calculate the passanger traffic submitted by DIAL and MIAL for fourth control period . divide the passenger traffic per manpoer cost anf compare it anf give us the rationale . Step 1: KPI Comparison. To begin, you will collect specific data from the DIAL Fourth Control Period Consultation Paper and DIAL Fourth Control Period Tariff Order, as well as the MIAL Fourth Control Period Consultation Paper and MIAL Fourth Control Period Tariff Order. From these documents, meticulously extract the manpower count, total passenger traffic, and total manpower expenditure for each fiscal year of their respective fourth control periods. With this comprehensive dataset, proceed to calculate two critical KPIs for both airports: manpower count per total passenger traffic and manpower expenditure per total passenger traffic. Once these KPIs are computed, compare them to BIAL's corresponding figures, assessing whether BIAL’s KPIs are higher, lower, or in line, while being careful to only compare data for years where the passenger traffic is similar to ensure the KPI comparison is accurate and meaningful. First, carefully examine BIAL's provided MDA document to identify the specific justifications for its manpower expense projections, including any explanations for variances from the prior control period. Next, to enhance this rationale, you will consult the detailed analyses and findings in the DIAL and MIAL Fourth Control Period Consultation Papers and Tariff Orders. Specifically, you will look for how these regulatory documents justify their own employee expense projections, such as by detailing factors like inflation, annual growth rates, and specific manpower growth factors tied to strategic operational expansions. Using these as a benchmark, you will then suggest improvements for BIAL's own justifications, for example, by recommending that BIAL provide a more granular breakdown of cost drivers, link employee growth to new projects or terminal expansions, or justify its average cost per employee based on specific salary benchmarks or industry-wide trends, ultimately making BIAL's rationale as transparent and well-supported as that of its peers."
            },
            "Utility Analysis": {
               "Analysis of electricity expenditure projection for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the electricity expenditure projected by BIAL: 1- Year on Year actual growth of power consumption cost for third control period  by BIAL in Utlities_cost_MDA document . 2-   Year on Year  actual power consumption by BIAL submitted in the Utlities_cost_MDA document. 3- Year on Year actual recoveries of power consumption by BIAL for third control period  in the Utlities_cost_MDA document. 4- Justification provided by BIAL for the power expense  and the variance of power expense with authority approved figures in third control period in the Utlities cost_MDA document. 5- Year on Year growth of actual power expense submitted by DIAL for true up of third control period in the fourth control period consultation paper. 6- Year on Year  growth of power consumption submitted by DIAL for third control period in the fourth control period consultation paper. 7- Year on Year actual recoveries from sub-concessionaries (%) submitted by DIAL for third control period in the fourth control period consultation paper. 8- Justification for actual power expense in third control period provided by DIAL and the variance with authority approved figures in fourth control period consultation paper. 9- Examination and rationale provided by authority on actual power cost and consumption submitted by DIAL for third control period in the fourth control period consultation paper.  10- Year on Year  Electricity cost(utility expenses) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 11- Year on Year  electricity  gross consumption(utlity expenses) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 12- Year on Year  recoveries of electricity consumption submitted by MIAL for the trueup of third control period in the MIAL fourth control period consultation paper. 8 Justification for actual electricity cost for the true up of third control period provided by MIAL in the MIALfourth control period consultation paper and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual Electricity cost and consumption submitted by MIAL true of third control period in the MIAL fourth control period consultation paper.15- Using the rationale extracted in steps 4, 8, 9,13 and 14 suggests how the rationale or justification provided by BIAL in the MDA document for electricity cost  for third control period can be enhanced. For every suggestion made, give specific reason why the suggestion was made using relevant references from DIAL and MIAL tariff orders or consultation papers. when asked about MIAL only give information relevant to MIAL not DIAL Strictly.",
               "Analysis of water expenditure projection for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the water expenditure projected by BIAL: 1-  actual portable and raw water cost by BIAL for trueup of third control period in Utlities_cost_MDA document . 2-year on Year raw and portable water  consumption by BIAL of true up for third control period in the Utilities cost_MDA document . 3- Year on Year actual recoveries of  water consumption by BIAL for the third control period in the Utlities_cost_MDA document . 4- Justification provided by BIAL for the water cost for third control period and the variance of water expense with authority approved figures in third control period in the Utlities_cost_MDA document. 5- Year on Year  water gross charge submitted by DIAL for true up of third control period in the DIAL fourth control period consultation paper. 6- Year on Year growth of water consumption submitted by DIAL for third control period in the DIAL fourth control period consultation paper. 7- Year on Year actual recoveries from sub- concessionaire submitted by DIAL for third control period in the DIAL fourth control period consultation paper. 8- Justification for actual  gross water charge  in third control period in the DIAL fourth control period consultation paper provided by DIAL and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual water gross charge and consumption submitted by DIAL for third control period in the DIAL fourth control period consultation paper.  10- Year on Year water expense(utility expenses) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 11- Year on Year water consumption(Kl) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 12- Year on Year  recoveries(kl) of water consumption submitted by MIAL for true up of the  third control period in the MIAL fourth control period consultation paper. 8- Justification for actual water gross amount for third control period in the MIAL fourth control period consultation paper provided by MIAL and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual water gross amount  and consumption submitted by MIAL for third control period in the MIAL fourth control period consultation paper.15- Using the rationale extracted in steps 4, 8, 9,13 and 14 suggest how the rationale or justification provided by BIAL in the MDA document for water expenditure for trueup of  third control period can be enhanced. For every suggestion made, give specific reason why the suggestion made using relevant references from DIAL and MIAL tariff orders or consultation papers.",
               "Analysis of KPI Computation for BIAL(Utility Expenditure)": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analyzing the KPI Computation. Calculate and compare the YoY change of power and electricity e expenses of DIAL and MIAL for true up of third control period p,first give what is total electricity expense submitted by DIAL for true up of third control period  , Electricty cost submitted by MIAL for true up of third control  period . after wards calculate the passanger traffic submitted by DIAL and MIAL for true up of third control period divide the passenger traffic per electricity cost  and compare it and give us the rationale ,Year on Year  water gross charge submitted by DIAL per passenger traffic submitted  for true up of third control period in the DIAL fourth control period consultation paper. Calculate and compare the YoY change of water a gross charge of DIAL and MIAL for true up of third control period p,first give what is total electricity expense submitted by DIAL for true up of third control period  , water cost submitted by MIAL for true up of third control  period . after wards calculate the passanger traffic submitted by DIAL and MIAL for true up of thord control perioddivide the passenger traffic per water cost and compare it and give us the rationale Step 1: KPI Comparison. To begin, you will collect specific data from the DIAL Fourth Control Period Consultation Paper and DIAL Fourth Control Period Tariff Order, as well as the MIAL Fourth Control Period Consultation Paper and MIAL Fourth Control Period Tariff Order. From these documents, meticulously extract the electricity consumption, water consumption, and total passenger traffic for each fiscal year of their respective fourth control periods. With this comprehensive dataset, proceed to calculate two critical KPIs for both airports: electricity consumption per total passenger traffic and water consumption per total passenger traffic. Once these KPIs are computed, compare them to BIAL's corresponding figures, assessing whether BIAL’s KPIs are higher, lower, or in line, while being careful to only compare data for years where the passenger traffic is similar to ensure the KPI comparison is accurate and meaningful. First, carefully examine BIAL's provided MDA document to identify the specific justifications for its utility expense projections, including any explanations for variances from the prior control period. Next, to enhance this rationale, you will consult the detailed analyses and findings in the DIAL and MIAL Fourth Control Period Consultation Papers and Tariff Orders. Specifically, you will look for how these regulatory documents justify their own utility expense projections, such as by detailing factors like energy efficiency initiatives, water conservation projects, infrastructure upgrades impacting consumption, or changes in operational scope. Using these as a benchmark, you will then suggest improvements for BIAL's own justifications, for example, by recommending that BIAL provide a more granular breakdown of consumption drivers, link utility usage to new terminal operations or technological advancements, or justify its per-passenger consumption figures based on industry best practices or environmental targets, ultimately making BIAL's rationale as transparent and well-supported as that of its peers."
            },
            "R&M Analysis": {
                "Analysis of repairs and maintenance expenditure for true up for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the repairs and maintenance expenditure projected by BIAL: 1- Year on Year actual growth of repairs and maintenance expenditure by BIAL for third control period in the MDA document 2- Year wise repairs and maintenance expenditure as a percentage of regulated asset base for BIAL for third control period in the MDA document. 3- Justification provided by BIAL for the repairs and maintenance expense for third control period and the variance of repairs and maintenance expense with authority approved figures in third control period in the MDA document. 4- Year on Year growth of actual repairs and maintenance expenditure submitted by DIAL for true up of third control period in the fourth control period consultation paper or tariff order. 5- Year wise repairs and maintenance expenditure as a percentage of regulated asset base for DIAL for third control period in the fourth control period consultation paper or tariff order. 6- Justification for actual repairs and maintenance expense in third control period provided by DIAL and the variance with authority approved figures for the third control period in fourth control period consultation paper or tariff order. 7- Examination and rationale provided by authority on actual repairs and maintenance cost submitted by DIAL for third control period in the fourth control period consultation paper or tariff order. 8- Year on Year growth of actual repairs and maintenance expenditure submitted by MIAL for true up of third control period in the fourth control period consultation paper or tariff order. 9- Justification for actual repairs and maintenance expense in third control period provided by MIAL and the variance with authority approved figures for the third control period in fourth control period consultation paper or tariff order. 10- Year wise repairs and maintenance expenditure as a percentage of regulated asset base for MIAL for third control period in the fourth control period consultation paper or tariff order. 11- Examination and rationale provided by authority on actual repairs and maintenance cost submitted by MIAL for third control period in the fourth control period consultation paper or tariff order. 12- Using the rationale extracted in steps 5, 6, 8, and 9 suggest how the rationale or justification provided by BIAL in the MDA document for repairs and maintenance expenditure for third control period can be enhanced. For every suggestion made, give specific reason why the suggestion is made using relevant references from DIAL and MIAL tariff orders or consultation papers",
                "Analysis of repairs and maintenance expenditure projection for BIAL for fourth control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the repairs and maintenance expenditure projected by BIAL: 1- Year on Year growth of repairs and maintenance expenditure projections by BIAL for fourth control period in the MDA document 2- Year wise repairs and maintenance expenditure projection as a percentage of regulated asset base for BIAL for fourth control period in the MDA document. 3- Justification provided by BIAL for the repairs and maintenance expense for fourth control period in the MDA document. 4- Year on Year growth of repairs and maintenance expenditure projections submitted by DIAL for fourth control period in the fourth control period consultation paper or tariff order. 5- Year wise repairs and maintenance expenditure projections as a percentage of regulated asset base for DIAL for fourth control period in the fourth control period consultation paper or tariff order. 6- Justification for repairs and maintenance expense projections in fourth control period provided by DIAL in fourth control period consultation paper or tariff order. 7- Examination and rationale provided by authority on repairs and maintenance expenditure projections submitted by DIAL for fourth control period in the fourth control period consultation paper or tariff order. 8- Year on Year growth of repairs and maintenance expenditure projections submitted by MIAL for fourth control period in the fourth control period consultation paper or tariff order. 9- Year wise repairs and maintenance expenditure projections as a percentage of regulated asset base for MIAL for fourth control period in the fourth control period consultation paper or tariff order. 10- Justification for repairs and maintenance expense projections in fourth control period provided by MIAL in fourth control period consultation paper or tariff order. 11- Examination and rationale provided by authority on repairs and maintenance expenditure projections submitted by MIAL for fourth control period in the fourth control period consultation paper or tariff order 12- Using the rationale extracted in steps 5, 6, 8, and 9 suggest how the rationale or justification provided by BIAL in the MDA document for repairs and maintenance expenditure for fourth control period can be enhanced. For every suggestion made, give specific reason why the suggestion is made using relevant references from DIAL and MIAL tariff orders or consultation papers"
            },
        }
        
        prompt_template = next((category[analysis_title] for category in analysis_prompts_config.values() if analysis_title in category), None)
        
        if not prompt_template:
            raise HTTPException(status_code=404, detail=f"Analysis title '{analysis_title}' not found.")
            
        full_prompt = prompt_template.format(document_summary=document_summary)
        
        # --- MODIFICATION: Use the unified Advanced RAG Pipeline ---
        logging.info(f"Starting Advanced RAG pipeline for document analysis: '{analysis_title}'")
        rag_result = await execute_advanced_rag_pipeline(user_question=full_prompt)
        analysis_answer = rag_result["answer"]
        # --- END MODIFICATION ---

        final_report = f"<h2>Analysis Report: {analysis_title}</h2><h3>Document Summary</h3><p>{document_summary}</p><hr /><h3>Analysis</h3>{analysis_answer}"
        return AnalysisResponse(report=final_report)
    except Exception as e:
        logging.error("AN ERROR OCCURRED IN /analyze-document", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An internal error occurred during analysis. Check server logs for details.")
    finally:
        duration = time.time() - start_time
        log_interaction(user=current_user.username, question=f"Initial Report: {analysis_title}", answer=final_report, duration=duration)

@app.post("/mda-chat", response_model=ConversationalChatResponse)
async def mda_chat(request: ChatRequest, current_user: User = Depends(auth_dependency)):
    return await handle_chat_request(request, current_user)

@app.post("/conversational-chat", response_model=ConversationalChatResponse)
async def conversational_chat(request: ChatRequest, current_user: User = Depends(auth_dependency)):
    return await handle_chat_request(request, current_user)

@app.post("/refine-report", response_model=RefineReportResponse)
async def refine_report(request: RefineReportRequest, current_user: User = Depends(auth_dependency)):
    if not clients.synthesis_openai_client:
        raise HTTPException(status_code=503, detail="OpenAI service is not configured.")
    try:
        refinement_prompt = f"""You are a report writing expert. Your task is to seamlessly integrate a new piece of information into an existing report. Do not simply append the new information. Instead, find the most relevant section in the 'ORIGINAL REPORT' and intelligently merge the 'NEW INFORMATION' into it. Rewrite paragraphs as needed to ensure the final report is coherent, clean, and well-integrated. Return ONLY the full, updated report text.
        **ORIGINAL REPORT:**
        ---
        {request.original_report}
        ---
        **NEW INFORMATION TO INTEGRATE:**
        ---
        {request.new_info}
        ---
        **FULL, REFINED, AND INTEGRATED REPORT:**"""
        
        response = clients.synthesis_openai_client.chat.completions.create(
            model=config.deployment_id, 
            messages=[{"role": "user", "content": refinement_prompt}], 
            temperature=0.2, 
            max_tokens= 16000
        )
        return RefineReportResponse(refined_report=response.choices[0].message.content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred during report refinement: {e}")

@app.post("/download-report")
async def download_report(request: DownloadRequest, background_tasks: BackgroundTasks, current_user: User = Depends(auth_dependency)):
    """
    MODIFIED: This endpoint now uses FileResponse instead of StreamingResponse to ensure
    the complete file is generated before download begins. It uses a background task
    to clean up the temporary file after sending.
    """
    try:
        document = Document()
        # Add default styles for lists if they don't exist
        styles = document.styles
        if 'List Bullet' not in styles:
            styles.add_style('List Bullet', 1)
        if 'List Number' not in styles:
            styles.add_style('List Number', 1)

        soup = BeautifulSoup(request.html_content, 'html.parser')

        # Use the new, comprehensive parsing function
        parse_html_to_docx(soup, document)

        # Save the document to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            document.save(tmp.name)
            temp_file_path = tmp.name
        
        # Add a background task to delete the file after the response has been sent
        background_tasks.add_task(os.remove, temp_file_path)

        # Return the complete file using FileResponse
        return FileResponse(
            path=temp_file_path,
            filename="BIAL_Analysis_Report.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logging.error("Error creating Word document", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error creating Word document: {e}")

@app.post("/feedback")
async def handle_feedback(request: FeedbackRequest, current_user: User = Depends(auth_dependency)):
    if not clients.mongo_client:
        raise HTTPException(status_code=503, detail="Database service is not configured.")
    try:
        feedback_collection = clients.mongo_client[config.cosmos_database_name][config.cosmos_feedback_collection]
        
        feedback_doc = {
            "timestamp": datetime.now(timezone.utc),
            "user": current_user.username,
            "question": request.question,
            "answer": request.answer,
            "feedback": request.feedback
        }
        
        feedback_collection.insert_one(feedback_doc)
        return {"status": "success", "message": "Feedback recorded."}
    except Exception as e:
        logging.error(f"Error recording feedback", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to record feedback: {e}")





